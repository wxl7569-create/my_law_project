# ══════════════════════════════════════════════════════════════
# 定义数据类，统一解析后的文档格式
# ══════════════════════════════════════════════════════════════

from dataclasses import dataclass

@dataclass
class ParsedDocument:
    filename: str          # 文件名
    content: str           # 统一输出的纯文本内容
    file_type: str         # 文件扩展名（如 'txt', 'docx'）
    page_count: int = 0    # 页数（非必须，留给PDF用）
    metadata: dict = None  # 预留扩展


# ══════════════════════════════════════════════════════════════
# 导入标准库与内部模块
# ══════════════════════════════════════════════════════════════

from abc import ABC, abstractmethod
import os
from .logger import LoggerManager

logger = LoggerManager.get_logger()


# ══════════════════════════════════════════════════════════════
# 基础解析器接口
# ══════════════════════════════════════════════════════════════

class BaseParser(ABC):
    @abstractmethod
    def parse(self, file_path: str) -> ParsedDocument:
        pass


# ══════════════════════════════════════════════════════════════
# 具体解析器实现
# ══════════════════════════════════════════════════════════════

class TxtParser(BaseParser):
    """解析 .txt 纯文本文件"""
    def parse(self, file_path: str) -> ParsedDocument:
        logger.info(f"TxtParser 开始解析: {file_path}")
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        logger.info(f"TxtParser 解析完成，内容长度={len(content)}")
        return ParsedDocument(
            filename=os.path.basename(file_path),
            content=content,
            file_type='txt'
        )


class DocxParser(BaseParser):
    """解析 .docx 文件"""
    def parse(self, file_path: str) -> ParsedDocument:
        logger.info(f"DocxParser 开始解析: {file_path}")
        from docx import Document
        doc = Document(file_path)
        content = "\n".join([para.text for para in doc.paragraphs if para.text])
        logger.info(f"DocxParser 解析完成，段落数={len(doc.paragraphs)}，内容长度={len(content)}")
        return ParsedDocument(
            filename=os.path.basename(file_path),
            content=content,
            file_type='docx'
        )


class DocParser(BaseParser):
    """针对老旧 .doc 格式（作为后备方案）"""
    def parse(self, file_path: str) -> ParsedDocument:
        logger.info(f"DocParser 开始解析: {file_path}")
        try:
            import textract  # 需安装 textract (依赖较多)
            content = textract.process(file_path).decode('utf-8')
            logger.info(f"DocParser 解析完成，内容长度={len(content)}")
        except ImportError:
            logger.warning("textract 未安装，暂不支持解析老旧 .doc 格式")
            content = "[错误] 暂不支持解析老旧 .doc 格式，请转为 .docx 或 .txt"
        return ParsedDocument(
            filename=os.path.basename(file_path),
            content=content,
            file_type='doc'
        )


class MdParser(TxtParser):
    """解析 .md Markdown 文件（按纯文本读取）"""
    def parse(self, file_path: str) -> ParsedDocument:
        doc = super().parse(file_path)
        doc.file_type = 'md'
        return doc


# ══════════════════════════════════════════════════════════════
# 解析器工厂
# ══════════════════════════════════════════════════════════════

class ParserFactory:
    """根据后缀名返回对应的解析器实例"""
    _parsers = {
        '.txt': TxtParser(),
        '.md': MdParser(),
        '.docx': DocxParser(),
        '.doc': DocParser(),
    }
    
    @classmethod
    def get_parser(cls, ext: str):
        logger.debug(f"ParserFactory 查找解析器: ext={ext}")
        parser = cls._parsers.get(ext.lower())
        if not parser:
            logger.error(f"不支持的文件格式: {ext}")
            raise ValueError(f"不支持的文件格式: {ext}")
        return parser


# ══════════════════════════════════════════════════════════════
# 对外统一入口（自动解析 + 注入Prompt）
# ══════════════════════════════════════════════════════════════

def parse_file_and_inject(file_path: str, custom_prompt_template: str = None) -> str:
    """
    解析文件并返回可以直接喂给LLM的Prompt字符串
    """
    logger.info(f"parse_file_and_inject 开始处理: {file_path}")

    # 1. 自动识别扩展名
    ext = os.path.splitext(file_path)[-1]
    
    # 2. 获取对应解析器并执行解析
    parser = ParserFactory.get_parser(ext)
    parsed_doc = parser.parse(file_path)
    
    # 3. 内容清洗（防止文本污染Prompt）
    clean_content = parsed_doc.content.strip()
    if not clean_content:
        logger.warning(f"文件内容为空或仅含空白字符: {file_path}")
        clean_content = "[警告] 解析出的内容为空，请检查文件是否损坏。"
    
    # 4. 默认Prompt模板（如果外部没传，就用内置的）
    if custom_prompt_template is None:
        prompt_template = """
你是一个专业的文档分析助手。请仔细阅读以下【文档内容】，并准备回答用户基于该文档提出的问题。

【文档名称】：{filename}
【文档类型】：{filetype}
【文档内容】：
{content}
"""
    else:
        prompt_template = custom_prompt_template
    
    # 5. 注入内容并返回
    final_prompt = prompt_template.format(
        filename=parsed_doc.filename,
        filetype=parsed_doc.file_type,
        content=clean_content
    )
    
    logger.info(f"parse_file_and_inject 完成，prompt 长度={len(final_prompt)}")
    return final_prompt


# ══════════════════════════════════════════════════════════════
# 测试入口
# ══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import tempfile     #创建临时文件和目录

    # 测试：传入文档 → 加载出合理的提示词
    print("word_reader 测试：文档 → Prompt")

    # 1. .txt 文件
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", encoding="utf-8", delete=False) as f:
        f.write("民法典第一百条测试内容")
        txt_path = f.name
    prompt = parse_file_and_inject(txt_path)
    print(f"\n[.txt] prompt 长度={len(prompt)}")
    print(prompt[:200])
    assert "民法典第一百条测试内容" in prompt, "内容未注入 prompt"
    assert "文档名称" in prompt, "默认模板缺失"
    os.unlink(txt_path)

    # 2. .docx 文件
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_paragraph("合同法测试段落")
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            docx_path = f.name
            doc.save(docx_path)
        prompt = parse_file_and_inject(docx_path)
        print(f"\n[.docx] prompt 长度={len(prompt)}")
        print(prompt[:200])
        assert "合同法测试段落" in prompt, "内容未注入 prompt"
        os.unlink(docx_path)
    except ImportError:
        print("\n[.docx] python-docx 未安装，跳过")

    # 3. 空文件 → 应包含警告
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", encoding="utf-8", delete=False) as f:
        f.write("   \n\t  ")
        empty_path = f.name
    prompt = parse_file_and_inject(empty_path)
    print(f"\n[空文件] prompt 长度={len(prompt)}")
    assert "警告" in prompt, "空文件未输出警告"
    os.unlink(empty_path)

    print("\n✓ 测试全部通过")
